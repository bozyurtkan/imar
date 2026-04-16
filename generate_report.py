from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx():
    doc = Document()

    # Title
    title = doc.add_heading('4 Nisan 2026 Tarihli Resmî Gazete Yönetmelik Analizi', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Konu: İzinsiz/Ruhsatsız Bağ Evi ve Bungalov Tarzı Yapıların Durumu ve Mal Sahiplerinin Atması Gereken Adımlar')

    # Section 1
    doc.add_heading('1. Genel Bakış', level=1)
    doc.add_paragraph(
        '4 Nisan 2026 tarihli Resmî Gazete\'de yayımlanan "Tarım Arazilerinin Korunması ve Kullanılması Hakkında Yönetmelik" '
        've "Arazi Kullanım Planlaması Uygulama Yönetmeliği" ile tarım arazilerindeki yapılaşma şartları yeniden düzenlenmiş '
        've izinsiz yapılara yönelik ciddi yaptırımlar getirilmiştir.'
    )

    # Section 2
    doc.add_heading('2. İzinsiz ve Ruhsatsız Yapılar İçin Yaptırımlar', level=1)
    p2 = doc.add_paragraph()
    p2.add_run('Hızlı Yıkım Süreci: ').bold = True
    p2.add_run(
        'İzinsiz yapıldığı tespit edilen yapılar için ilgili belediye veya il özel idaresine bildirim yapılır. '
        'Bu kurumların 1 ay içinde yıkımı gerçekleştirmesi zorunludur. Aksi takdirde Tarım ve Orman Bakanlığı yıkımı bizzat '
        'yapacak ve masrafları %100 fazlasıyla ilgili yerel yönetimden tahsil edecektir.'
    )
    
    p3 = doc.add_paragraph()
    p3.add_run('Eski Haline Getirme Bedeli: ').bold = True
    p3.add_run(
        'Yıkım sonrası arazinin tarımsal vasfına döndürülmesi için yapılan tüm masraflar mal sahibinden tahsil edilir.'
    )

    p4 = doc.add_paragraph()
    p4.add_run('Adli Soruşturma: ').bold = True
    p4.add_run(
        'Özellikle "hobi bahçesi" adı altında tarım arazilerinin bölünmesine ve yapılaşmasına aracılık edenler '
        'hakkında Cumhuriyet Başsavcılığına suç duyurusunda bulunulacaktır.'
    )

    # Section 3
    doc.add_heading('3. Bağ Evi Kriterleri (Yasal Şartlar)', level=1)
    doc.add_paragraph('Bir yapının "bağ evi" olarak yasal statü kazanabilmesi için şu şartları sağlaması gerekmektedir:')
    doc.add_paragraph('• Arazi Büyüklüğü: Mutlak/Özel ürün arazilerinde min. 5000 m², dikili tarım arazilerinde min. 3000 m².', style='List Bullet')
    doc.add_paragraph('• İnşaat Alanı: Brüt 100 m²’yi geçemez (müştemilat dahil).', style='List Bullet')
    doc.add_paragraph('• Kat Sınırı: Bodrum hariç tek katlı.', style='List Bullet')
    doc.add_paragraph('• İzin/Ruhsat: Tarımsal amaçlı yapı izni ve yapı ruhsatı alınması şarttır.', style='List Bullet')

    # Section 4
    doc.add_heading('4. Bungalov ve Tiny House Durumu', level=1)
    doc.add_paragraph(
        'Yönetmelik, tarım arazisi üzerine izinsiz kondurulan veya sabitlenen bungalov ve tekerlekli yapıları (tiny house) '
        'doğrudan "izinsiz yapı" statüsünde değerlendirmektedir. Bu yapılar eğer tarımsal amaçlı yapı iznine sahip değilse, '
        'yıkım ve para cezası hükümlerine tabidir.'
    )

    # Section 5
    doc.add_heading('5. Mal Sahiplerinin Yapması Gerekenler', level=1)
    doc.add_paragraph('1. Statü Kontrolü: Arazinin tapu kaydını ve imar durumunu "parsel.tkgm.gov.tr" veya ilgili ilçe tarım müdürlüğünden kontrol edin.', style='List Number')
    doc.add_paragraph('2. İskan ve Ruhsat Araştırması: Mevcut yapınız için alınmış bir yapı kayıt belgesi veya ruhsat olup olmadığını teyit edin.', style='List Number')
    doc.add_paragraph('3. Yasal Başvuru: Eğer arazi büyüklüğü kriterleri (5000/3000 m²) sağlanıyorsa, yapıyı yasal hale getirmek için "Tarımsal Amaçlı Yapı İzni" başvurusu imkanlarını hukuk müşavirinizle değerlendirin.', style='List Number')
    doc.add_paragraph('4. Gönüllü Tahliye/Eski Haline Getirme: Yasal kriterleri sağlamayan ve yıkım kararı kesinleşebilecek yapılar için, yıkım masraflarının katlanarak size rücu etmesini önlemek adına yapıyı kaldırmayı değerlendirin.', style='List Number')

    # Conclusion
    doc.add_heading('6. Sonuç', level=1)
    doc.add_paragraph(
        'Yeni yönetmelik tarım arazilerini koruma konusunda oldukça tavizsiz bir tutum sergilemektedir. '
        'Özellikle "hobi bahçesi" tarzı oluşumlarda yer alan ruhsatsız yapıların korunması hukuken mümkün görünmemektedir. '
        'Mal sahiplerinin ivedilikle arazilerinin hukuki durumunu netleştirmeleri önerilir.'
    )

    # Footer
    doc.add_paragraph('\nNot: Bu rapor 4 Nisan 2026 tarihli Resmî Gazete (Sayı: 33000+ serisi) verilerine dayanarak hazırlanmıştır.')

    doc.save('Mevzuat_Analiz_Raporu.docx')
    print("Mevzuat_Analiz_Raporu.docx başarıyla oluşturuldu.")

if __name__ == "__main__":
    create_docx()
